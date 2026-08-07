"""文档解析测试。"""

from docx import Document

from backend.rag.parser import (
    _parse_docx_sync,
    _parse_md_sync,
    _parse_txt_sync,
    parse_md,
    parse_txt,
)


def test_parse_docx_paragraphs_unique_page(tmp_path):
    """回归：多段 DOCX 每段必须有唯一 page，避免 chunk_id 碰撞。"""
    d = Document()
    d.add_paragraph("第一段：水利工程是治理水患的工程。")
    d.add_paragraph("第二段：水库调度统筹防洪与兴利。")
    d.add_paragraph("第三段：防汛预案需定期演练。")
    fp = tmp_path / "multi.docx"
    d.save(str(fp))

    pages = _parse_docx_sync(fp)
    assert len(pages) == 3
    assert [p["page"] for p in pages] == [1, 2, 3]
    assert all(p["content"].strip() for p in pages)


def test_parse_docx_skips_empty_paragraphs(tmp_path):
    d = Document()
    d.add_paragraph("")
    d.add_paragraph("有内容的一段")
    d.add_paragraph("   ")
    fp = tmp_path / "empty.docx"
    d.save(str(fp))

    pages = _parse_docx_sync(fp)
    assert len(pages) == 1  # 只保留非空段落
    assert pages[0]["page"] >= 1  # 页码可能因空段落跳过，但必须为正整数


def test_parse_txt_pages_by_50_lines(tmp_path):
    """TXT 按每 50 行打一页，页码连续递增。"""
    fp = tmp_path / "a.txt"
    fp.write_text("\n".join(f"line {i}" for i in range(120)), encoding="utf-8")

    pages = _parse_txt_sync(fp)
    assert len(pages) == 3  # 50 + 50 + 20
    assert [p["page"] for p in pages] == [1, 2, 3]
    assert all(p["content"].strip() for p in pages)


def test_parse_txt_empty_falls_back_single_page(tmp_path):
    """TXT 全空白：回退为单页（空内容让上游 chunk 过滤丢弃）。"""
    fp = tmp_path / "empty.txt"
    fp.write_text("   \n  \n", encoding="utf-8")

    pages = _parse_txt_sync(fp)
    assert len(pages) == 1
    assert pages[0]["page"] == 1


def test_parse_md_single_page(tmp_path):
    """MD 当作单页大文本。"""
    fp = tmp_path / "b.md"
    text = "# 标题\n\n正文第一段。\n\n正文第二段。"
    fp.write_text(text, encoding="utf-8")

    pages = _parse_md_sync(fp)
    assert len(pages) == 1
    assert pages[0]["page"] == 1
    assert pages[0]["content"] == text


async def test_parse_txt_md_async_wrappers(tmp_path):
    """G10.19：TXT/MD 解析走异步包装（读盘在 executor，不阻塞事件循环）。"""
    t = tmp_path / "c.txt"
    t.write_text("\n".join(f"row {i}" for i in range(110)), encoding="utf-8")
    assert len(await parse_txt(t)) == 3

    m = tmp_path / "c.md"
    m.write_text("md 内容", encoding="utf-8")
    pages = await parse_md(m)
    assert len(pages) == 1
    assert pages[0]["content"] == "md 内容"
